"""Generate the SPI Methodology Technical Note as a Word document."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width    = Inches(8.5)
section.page_height   = Inches(11)
section.left_margin   = Inches(1.1)
section.right_margin  = Inches(1.1)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)


def set_color(run, hex_str):
    h = hex_str.lstrip("#")
    run.font.color.rgb = RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def h1(text, color="#003d82"):
    p = doc.add_heading(text, level=1)
    if p.runs:
        p.runs[0].font.size = Pt(18)
        set_color(p.runs[0], color)
    return p


def h2(text, color="#003d82"):
    p = doc.add_heading(text, level=2)
    if p.runs:
        p.runs[0].font.size = Pt(14)
        set_color(p.runs[0], color)
    return p


def h3(text, color="#0070c0"):
    p = doc.add_heading(text, level=3)
    if p.runs:
        p.runs[0].font.size = Pt(12)
        set_color(p.runs[0], color)
    return p


def body(text, size=11, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        set_color(run, color)
    return p


def bullet(text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def formula_box(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.bold = True
    set_color(run, "#003d82")
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "EEF4FF")
    pPr.append(shd)
    return p


def note(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic = True
    set_color(run, "#6b7280")
    return p


def hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C5D5E8")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def tbl_header(table, headers):
    cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cells[i].text = h
        if cells[i].paragraphs[0].runs:
            cells[i].paragraphs[0].runs[0].bold = True


def tbl_row(table, row_idx, values):
    cells = table.rows[row_idx].cells
    for i, v in enumerate(values):
        cells[i].text = v


# ── COVER / TITLE ─────────────────────────────────────────────────────────────
doc.add_paragraph()
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("Surveillance Performance Index (SPI)")
r.font.size = Pt(22)
r.bold = True
set_color(r, "#003d82")

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sp.add_run("Methodology, Formula & Worked Examples")
r2.font.size = Pt(14)
set_color(r2, "#4a90e2")

op = doc.add_paragraph()
op.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = op.add_run("NCD Population-based Surveillance Intelligence Platform  |  WHO African Region")
r3.font.size = Pt(11)
set_color(r3, "#6b7280")

doc.add_paragraph()
hr()
doc.add_paragraph()

# ── 1. PURPOSE ────────────────────────────────────────────────────────────────
h2("1.  Purpose and Scope")
body(
    "The Surveillance Performance Index (SPI) is a composite score ranging from 0 to 100 "
    "that quantifies how well a country's NCD population-based surveillance system performs "
    "across three orthogonal dimensions: Breadth, Currency, and Regularity."
)
doc.add_paragraph()
body("The index is designed for:")
bullet("Relative ranking of WHO AFRO Member States and territories (N = 48).")
bullet("Tracking improvement or decline in surveillance capacity over time.")
bullet("Informing strategic prioritisation and resource allocation decisions.")
doc.add_paragraph()
note(
    "Important: SPI scores reflect surveillance system performance, not the health status "
    "of the population. A high SPI means a country has robust, current, and regular data; "
    "it does not imply low NCD burden."
)

# ── 2. DATA INPUTS ─────────────────────────────────────────────────────────────
doc.add_paragraph()
h2("2.  Input Data")
body(
    "The SPI is computed from the WHO NCD Country Tracking dataset for the WHO African Region. "
    "For each country, the model uses:"
)
bullet("The list of completed STEPS, GYTS, GSHS, GATS, and GSHPP surveys (year + status).")
bullet(
    'Only records with status = "Completed (At least data analysis completed)" are counted.'
)
bullet('"In Progress" and "Not Usable" records are excluded from SPI computation.')
bullet("Gaps are measured in years relative to the reference year (2026).")

# ── 3. DIMENSIONS ─────────────────────────────────────────────────────────────
doc.add_paragraph()
h2("3.  The Three SPI Dimensions")

# 3a Breadth
doc.add_paragraph()
h3("3a.  Breadth Score (BS)  —  Institutional Range")
body(
    "Question answered: How many of the five NCD surveillance instruments has this country ever used?",
    italic=True,
)
doc.add_paragraph()
body(
    "For each of the 5 survey instruments (STEPS, GYTS, GSHS, GATS, GSHPP), the model checks "
    "whether the country has at least one completed record. BS is the count of instruments ever "
    "completed divided by 5:"
)
doc.add_paragraph()
formula_box("BS  =  (Number of instrument types completed \u2265 1 time)  /  5")
doc.add_paragraph()
body("BS is discrete and takes values in: {0, 0.2, 0.4, 0.6, 0.8, 1.0}")
body(
    "It is permanent and structural \u2014 past completions are never forgotten. "
    "It does not decay over time."
)

doc.add_paragraph()
body("Example \u2014 Breadth Score", bold=True)
t1 = doc.add_table(rows=4, cols=3)
t1.style = "Table Grid"
tbl_header(t1, ["Country", "Instruments completed", "BS"])
tbl_row(t1, 1, ["Country A", "STEPS, GYTS, GSHS  (3 of 5)", "0.60"])
tbl_row(t1, 2, ["Country B", "STEPS only  (1 of 5)", "0.20"])
tbl_row(t1, 3, ["Country C", "None  (0 of 5)", "0.00"])

# 3b Currency
doc.add_paragraph()
doc.add_paragraph()
h3("3b.  Currency Score (CS)  —  Evidence Recency")
body(
    "Question answered: How current is the evidence across the full surveillance portfolio?",
    italic=True,
)
doc.add_paragraph()
body(
    "CS uses exponential decay to score evidence freshness. For each of the 5 instruments, "
    "a decay value is calculated based on how many years (g) have elapsed since the last "
    "completed survey:"
)
doc.add_paragraph()
formula_box("Decay value for one instrument  =  exp(\u2212 ln(2)/7  \u00d7  g)")
doc.add_paragraph()
body(
    "Where g = years since the last completed survey (relative to 2026). The half-life is "
    "7 years, meaning a survey completed exactly 7 years ago contributes half the score of "
    "a survey completed today."
)
body("If an instrument has never been completed, its decay value is 0.")
doc.add_paragraph()
formula_box("CS  =  mean of decay values across all 5 instruments")
doc.add_paragraph()
body("Example \u2014 Currency Score for Country A", bold=True)
t2 = doc.add_table(rows=6, cols=3)
t2.style = "Table Grid"
tbl_header(t2, ["Instrument", "Last completed", "Decay value"])
tbl_row(t2, 1, ["STEPS",  "2022  (4 years ago)",  "exp(\u22120.099 \u00d7 4) = 0.671"])
tbl_row(t2, 2, ["GYTS",   "2014  (12 years ago)", "exp(\u22120.099 \u00d7 12) = 0.301"])
tbl_row(t2, 3, ["GSHS",   "2019  (7 years ago)",  "exp(\u22120.099 \u00d7 7) = 0.500"])
tbl_row(t2, 4, ["GATS",   "Never",                "0.000"])
tbl_row(t2, 5, ["GSHPP",  "Never",                "0.000"])
doc.add_paragraph()
body("CS = (0.671 + 0.301 + 0.500 + 0.000 + 0.000) / 5 = 1.472 / 5 = 0.294")

# 3c Regularity
doc.add_paragraph()
h3("3c.  Coverage-Adjusted Regularity (CA-ARI)  —  Cycling Behaviour")
body(
    "Question answered: Does the country survey on the WHO-recommended 5-year schedule, "
    "for instruments it has completed more than once?",
    italic=True,
)
doc.add_paragraph()
body(
    "For each instrument where the country has >= 2 completed rounds, the average interval "
    "between consecutive surveys is computed. A regularity score (d) is assigned:"
)
doc.add_paragraph()
formula_box("d  =  min(1,  5  /  average interval in years)")
doc.add_paragraph()
body(
    "d = 1.0 means the country cycles every 5 years or less (on schedule or better). "
    "d < 1.0 means the average gap is longer than 5 years."
)
doc.add_paragraph()
body(
    "ARI (Average Regularity Index) is the mean of d across all assessable instruments "
    "(those with >= 2 rounds). CA-ARI adjusts for countries that can only be assessed on "
    "a subset of the 5 instruments:"
)
doc.add_paragraph()
formula_box("ARI     =  mean(d)  across all assessable instruments")
formula_box("CA-ARI  =  ARI  \u00d7  (number of assessable instruments  /  5)")
doc.add_paragraph()
body(
    "If no instrument has been completed >= 2 times, CA-ARI = 0. This is not a penalty \u2014 "
    "it simply reflects that cycling behaviour cannot yet be evaluated."
)
doc.add_paragraph()
body("Example \u2014 CA-ARI for Country B", bold=True)
body("Country B completed STEPS in 2008, 2014, and 2022; and GYTS in 2010 and 2018.")
t3 = doc.add_table(rows=3, cols=4)
t3.style = "Table Grid"
tbl_header(t3, ["Instrument", "Completed rounds", "Avg interval", "d"])
tbl_row(t3, 1, ["STEPS", "2008, 2014, 2022", "(6+8)/2 = 7.0 yr", "min(1, 5/7.0) = 0.714"])
tbl_row(t3, 2, ["GYTS",  "2010, 2018",       "8.0 yr",            "min(1, 5/8.0) = 0.625"])
doc.add_paragraph()
body("ARI = (0.714 + 0.625) / 2 = 0.670")
body("CA-ARI = 0.670 \u00d7 (2 / 5) = 0.268")

# ── 4. COMBINING ─────────────────────────────────────────────────────────────
doc.add_paragraph()
h2("4.  Combining the Three Dimensions into SPI")
body(
    "The three dimension scores are averaged and scaled to a 0\u2013100 range:"
)
doc.add_paragraph()
formula_box("SPI  =  100  \u00d7  (BS + CS + CA-ARI) / 3")
doc.add_paragraph()
body(
    "Each dimension is in the range [0, 1] before scaling, so SPI \u2208 [0, 100]. "
    "The equal weighting of 1/3 each reflects the conceptual independence of breadth, "
    "currency, and regularity as distinct aspects of surveillance quality."
)

# ── 5. FULL WORKED EXAMPLE ────────────────────────────────────────────────────
doc.add_paragraph()
h2("5.  Full Worked Example \u2014 Ethiopia (4 STEPS Rounds)")
body(
    "Ethiopia has completed STEPS in 2003, 2006, 2015, and 2024. "
    "It has not completed GYTS, GSHS, GATS, or GSHPP."
)
doc.add_paragraph()

h3("Step 1 \u2014 Breadth Score")
body("Instruments ever completed: STEPS (1 of 5)")
body("BS = 1 / 5 = 0.200", bold=True)

h3("Step 2 \u2014 Currency Score  (reference year = 2026)")
t4 = doc.add_table(rows=6, cols=3)
t4.style = "Table Grid"
tbl_header(t4, ["Instrument", "Last completed / gap", "Decay value"])
tbl_row(t4, 1, ["STEPS",  "2024  (2 years ago)", "exp(\u22120.099 \u00d7 2) = 0.820"])
tbl_row(t4, 2, ["GYTS",   "Never",               "0.000"])
tbl_row(t4, 3, ["GSHS",   "Never",               "0.000"])
tbl_row(t4, 4, ["GATS",   "Never",               "0.000"])
tbl_row(t4, 5, ["GSHPP",  "Never",               "0.000"])
doc.add_paragraph()
body("CS = (0.820 + 0 + 0 + 0 + 0) / 5 = 0.164", bold=True)

h3("Step 3 \u2014 CA-ARI")
body("STEPS has 4 rounds: 2003, 2006, 2015, 2024.")
body("Intervals: 3, 9, 9 years  \u2192  average = (3+9+9)/3 = 7.0 years")
body("d(STEPS) = min(1, 5/7.0) = 0.714")
body("Only STEPS is assessable (other 4 instruments have < 2 rounds).")
body("ARI = 0.714")
body("CA-ARI = 0.714 \u00d7 (1 / 5) = 0.143", bold=True)

h3("Step 4 \u2014 Final SPI")
formula_box("SPI  =  100 \u00d7 (0.200 + 0.164 + 0.143) / 3  =  100 \u00d7 0.507 / 3  =  16.9")
doc.add_paragraph()
note(
    "Ethiopia's SPI of 16.9 places it in the Critical tier (< 30). This reflects a surveillance "
    "system with narrow instrument coverage (STEPS only) and irregular cycling \u2014 not a "
    "reflection of health outcomes. The 2024 STEPS round is recent, which boosts Currency; "
    "if Ethiopia expands to other instruments and maintains the 5-year cycle, its SPI will rise significantly."
)

# ── 6. COMPARISON OF TWO COUNTRIES ────────────────────────────────────────────
doc.add_paragraph()
h2("6.  Side-by-Side Comparison: Two Hypothetical Countries")
t5 = doc.add_table(rows=6, cols=3)
t5.style = "Table Grid"
tbl_header(t5, ["Dimension", "Country X", "Country Y"])
tbl_row(t5, 1, [
    "Survey history",
    "STEPS 2018 only",
    "STEPS 2010+2020, GYTS 2012+2019, GSHS 2015",
])
tbl_row(t5, 2, ["BS  (Breadth)", "1/5 = 0.20", "3/5 = 0.60"])
tbl_row(t5, 3, [
    "CS  (Currency)",
    "exp(\u22120.099\u00d78)/5 = 0.201/5 = 0.040",
    "Mean of 5 decay values = ~0.29",
])
tbl_row(t5, 4, [
    "CA-ARI  (Regularity)",
    "0.00  (only 1 STEPS round)",
    "~0.17  (2 assessable instruments)",
])
tbl_row(t5, 5, [
    "SPI",
    "100\u00d7(0.20+0.04+0.00)/3 = 8.0",
    "100\u00d7(0.60+0.29+0.17)/3 = 35.3",
])
doc.add_paragraph()
body(
    "Country X has only one instrument and it is 8 years old. Country Y is broader, its "
    "evidence is more current, and it has established cycling on two instruments. "
    "Y scores more than four times higher despite having a similar most-recent survey gap on STEPS."
)

# ── 7. TIER CLASSIFICATION ────────────────────────────────────────────────────
doc.add_paragraph()
h2("7.  Tier Classification")
t6 = doc.add_table(rows=5, cols=3)
t6.style = "Table Grid"
tbl_header(t6, ["Tier", "SPI Range", "Interpretation"])
tbl_row(t6, 1, [
    "Strong",     "SPI \u2265 75",
    "Broad, current, and regular surveillance. Evidence is policy-ready across all domains.",
])
tbl_row(t6, 2, [
    "Advancing",  "50 \u2013 74",
    "Solid foundation; gaps in one or two dimensions. Improvement within reach.",
])
tbl_row(t6, 3, [
    "Developing", "30 \u2013 49",
    "Significant weaknesses. Targeted support needed to improve breadth, currency, or regularity.",
])
tbl_row(t6, 4, [
    "Critical",   "SPI < 30",
    "Severe system-wide deficits. High priority for reinvestment and technical assistance.",
])

# ── 8. DESIGN PRINCIPLES ─────────────────────────────────────────────────────
doc.add_paragraph()
h2("8.  Key Design Principles")
bullet(
    "Orthogonality \u2014 The three dimensions measure conceptually independent aspects. "
    "Improving breadth requires a qualitatively different intervention than improving currency or regularity."
)
bullet(
    "No double-penalisation \u2014 A country that has never conducted GATS is not penalised twice "
    "(once in Breadth for never completing it, and again in Currency for having an old result). "
    "If an instrument was never completed, it contributes 0 to Currency and is excluded from Regularity."
)
bullet(
    "Universal formula \u2014 A single formula handles all cases. There is no special formula for "
    "nascent surveillance systems. A country with zero completions scores SPI = 0."
)
bullet(
    "Proportional penalisation for partial assessability \u2014 CA-ARI scales down proportionally "
    "when only some instruments are assessable, avoiding overstatement of regularity for countries "
    "with narrow coverage."
)
bullet(
    "Continuous currency \u2014 The exponential decay produces a continuous score. "
    "There are no threshold artefacts."
)

# ── 9. LIMITATIONS ────────────────────────────────────────────────────────────
doc.add_paragraph()
h2("9.  Known Limitations")
bullet(
    "SPI does not measure the quality or representativeness of individual surveys \u2014 "
    "only their existence, recency, and scheduling regularity."
)
bullet(
    "The 5-year cycle standard and 7-year currency half-life are normative choices. "
    "Alternative values would produce different scores, though rankings would be broadly stable."
)
bullet(
    "Countries with a single completed round cannot be assessed on Regularity. "
    "Their SPI is computed without that dimension, which may understate their true performance "
    "if they have recently established a surveillance system."
)
bullet(
    "The gap is calculated relative to a fixed reference year (2026). "
    "Re-running the pipeline in a future year will automatically update all scores."
)

doc.add_paragraph()
hr()
note(
    "Prepared by: NCD Population-based Surveillance Intelligence Platform \u2014 "
    "WHO African Region, Disease Prevention and Control (DPC) Cluster.  |  Reference year: 2026."
)

# ── Save ──────────────────────────────────────────────────────────────────────
out = "output/SPI_Methodology_Technical_Note.docx"
doc.save(out)
print("Saved: " + out)
